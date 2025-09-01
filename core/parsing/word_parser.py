from __future__ import annotations
from typing import Dict, Any, List
from io import BufferedIOBase
from docx import Document


def _join_clean(lines: List[str]) -> str:
    # Keep line breaks but drop empty-only lines
    return "\n".join([ln.strip() for ln in lines if str(ln).strip()])


def _extract_paragraphs(doc: Document) -> str:
    return _join_clean([p.text for p in doc.paragraphs])


def _extract_tables(doc: Document) -> List[List[List[str]]]:
    tables: List[List[List[str]]] = []
    for t in doc.tables:
        rows = []
        for r in t.rows:
            rows.append([c.text.strip() for c in r.cells])
        tables.append(rows)
    return tables


def _extract_headers_footers(doc: Document) -> Dict[str, str]:
    headers: List[str] = []
    footers: List[str] = []
    try:
        for sec in doc.sections:
            try:
                headers.extend([p.text for p in sec.header.paragraphs])
            except Exception:
                pass
            try:
                footers.extend([p.text for p in sec.footer.paragraphs])
            except Exception:
                pass
    except Exception:
        # some docs may not expose sections cleanly
        pass
    return {
        "header_text": _join_clean(headers),
        "footer_text": _join_clean(footers),
    }


def parse_word(file) -> Dict[str, Any]:
    """
    Minimal but useful DOCX parser.
    Returns:
      {
        "type": "docx",
        "text": "<all paragraph text>",
        "tables": [ [ [cell, ...], ... ], ... ],
        "headers": "<header text>",
        "footers": "<footer text>",
        "meta": { "paragraph_count": int, "table_count": int }
      }
    """
    # python-docx accepts a file-like object or a path.
    # Streamlit's UploadedFile supports .read()/.seek() so this works directly.
    doc = Document(file)

    text = _extract_paragraphs(doc)
    tables = _extract_tables(doc)
    hf = _extract_headers_footers(doc)

    item: Dict[str, Any] = {
        "type": "docx",
        "text": text,
        "tables": tables,
        "headers": hf.get("header_text", ""),
        "footers": hf.get("footer_text", ""),
        "meta": {
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(tables),
        },
    }
    return item
